"""
GhostResume.ai — FastAPI Backend
Production API that wraps the resume bot pipeline.
"""
import os
import re
import json
import uuid
import tempfile
from datetime import datetime
from pathlib import Path

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel
from typing import Optional

import httpx
import pdfplumber
from docx import Document as DocxDocument

app = FastAPI(title="GhostResume.ai", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Lock down in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Config
ANTHROPIC_KEY = os.getenv("ANTHROPIC_API_KEY", "")
PERPLEXITY_KEY = os.getenv("PERPLEXITY_API_KEY", "")
CLAUDE_MODEL = os.getenv("CLAUDE_MODEL", "claude-sonnet-4-20250514")
OUTPUT_DIR = Path("./output")
OUTPUT_DIR.mkdir(exist_ok=True)


# ============================================================
# HELPERS
# ============================================================
async def call_claude(system_prompt: str, user_message: str, use_search: bool = False) -> str:
    """Call Claude API and return text response."""
    payload = {
        "model": CLAUDE_MODEL,
        "max_tokens": 8000,
        "system": system_prompt,
        "messages": [{"role": "user", "content": user_message}],
    }
    if use_search:
        payload["tools"] = [{"type": "web_search_20250305", "name": "web_search"}]

    headers = {
        "x-api-key": ANTHROPIC_KEY,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }

    async with httpx.AsyncClient(timeout=120) as client:
        response = await client.post(
            "https://api.anthropic.com/v1/messages",
            headers=headers,
            json=payload,
        )
        response.raise_for_status()
        data = response.json()
        text_blocks = [b["text"] for b in data.get("content", []) if b.get("type") == "text"]
        return "\n".join(text_blocks)


async def call_perplexity(query: str) -> str:
    """Call Perplexity API for company research."""
    if not PERPLEXITY_KEY:
        return ""

    headers = {
        "Authorization": f"Bearer {PERPLEXITY_KEY}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": "sonar-pro",
        "messages": [
            {"role": "system", "content": "You are a company research analyst. Provide factual, current information."},
            {"role": "user", "content": query},
        ],
        "max_tokens": 2048,
    }

    async with httpx.AsyncClient(timeout=60) as client:
        try:
            response = await client.post(
                "https://api.perplexity.ai/chat/completions",
                headers=headers,
                json=payload,
            )
            response.raise_for_status()
            data = response.json()
            return data["choices"][0]["message"]["content"]
        except Exception:
            return ""


def parse_json_response(raw: str) -> dict:
    """Extract JSON from Claude response."""
    cleaned = raw.strip()
    if "```json" in cleaned:
        cleaned = cleaned.split("```json")[1].split("```")[0]
    elif "```" in cleaned:
        cleaned = cleaned.split("```")[1].split("```")[0]
    brace_start = cleaned.find("{")
    if brace_start > 0:
        cleaned = cleaned[brace_start:]
    brace_end = cleaned.rfind("}")
    if brace_end > -1:
        cleaned = cleaned[:brace_end + 1]
    return json.loads(cleaned.strip())


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file."""
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    doc = DocxDocument(file_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])


# ============================================================
# MODELS
# ============================================================
class PipelineRequest(BaseModel):
    resume_text: str
    job_input: str
    job_is_url: bool = False
    voice_mode: str = "match"  # "match" or "professional"


class PipelineResponse(BaseModel):
    session_id: str
    company: str
    role: str
    location: str
    ceo_pain: str
    ats_score: int
    viability_score: int
    recommendation: str
    tailored_resume: dict
    cover_letter: dict
    gap_report: dict
    interview_prep: dict
    red_flags: list
    green_flags: list


# ============================================================
# ENDPOINTS
# ============================================================
@app.get("/")
async def root():
    return {"name": "GhostResume.ai", "status": "running", "version": "1.0.0"}


@app.post("/api/upload-resume")
async def upload_resume(file: UploadFile = File(...)):
    """Upload and parse a resume file (PDF or DOCX) into text."""
    if not file.filename:
        raise HTTPException(400, "No file provided")

    ext = file.filename.lower().split(".")[-1]
    if ext not in ("pdf", "docx", "doc", "txt"):
        raise HTTPException(400, "Unsupported file type. Use PDF, DOCX, or TXT.")

    # Save temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name

    try:
        if ext == "pdf":
            text = extract_text_from_pdf(tmp_path)
        elif ext in ("docx", "doc"):
            text = extract_text_from_docx(tmp_path)
        else:
            text = content.decode("utf-8", errors="replace")
    finally:
        os.unlink(tmp_path)

    if not text.strip():
        raise HTTPException(400, "Could not extract text from file. Try pasting your resume text directly.")

    return {"text": text, "filename": file.filename, "char_count": len(text)}


@app.post("/api/run-pipeline")
async def run_pipeline(request: PipelineRequest):
    """Run the full GhostResume pipeline."""
    if not ANTHROPIC_KEY:
        raise HTTPException(500, "Anthropic API key not configured")

    session_id = str(uuid.uuid4())[:8]

    # Fetch job posting if URL
    job_text = request.job_input
    if request.job_is_url:
        try:
            async with httpx.AsyncClient(timeout=30, follow_redirects=True) as client:
                resp = await client.get(request.job_input, headers={
                    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7)"
                })
                text = re.sub(r'<script[^>]*>.*?</script>', '', resp.text, flags=re.DOTALL)
                text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
                text = re.sub(r'<[^>]+>', ' ', text)
                job_text = re.sub(r'\s+', ' ', text).strip()[:15000]
        except Exception as e:
            raise HTTPException(400, f"Could not fetch job posting URL: {e}")

    # Company research via Perplexity (parallel would be nice but keeping it simple)
    # We'll extract company name first from the job text
    company_research = ""

    # Main pipeline call to Claude
    voice_instruction = ""
    voice_consistency = ""
    if request.voice_mode == "match":
        voice_instruction = "2. VOICE PRINT: Analyze HOW the candidate naturally writes — sentence length, vocabulary level, formality, action verb preferences, whether they lead with results or context, use of jargon vs plain language. The tailored output must sound like a BETTER version of them, not a different person."
        voice_consistency = "- VOICE CONSISTENCY: Every bullet, the summary, and the cover letter must match the voice_print. If the candidate writes in short punchy sentences, the output uses short punchy sentences. If they use technical jargon naturally, keep it. The output should feel like the candidate on their best day, not like a different person wrote it."
    else:
        voice_instruction = "2. VOICE MODE: User selected PROFESSIONAL TONE. Ignore the candidate's original writing style. Write all output in a polished, professional, formal voice with strong action verbs, concise sentences, and industry-standard resume language. Still extract the voice_print for reference but do NOT match it — override it with professional tone."
        voice_consistency = "- VOICE OVERRIDE: Professional tone selected. All output should use polished, formal, recruiter-approved language regardless of how the original resume was written. Set voice_print.formality to 'professional' and voice_print.personality_notes to 'Professional tone selected by user'."

    system_prompt = (
        "You are GhostResume.ai — an expert resume tailoring engine that thinks like a recruiter.\n\n"
        "You use the 'Ghost Resume' methodology:\n"
        "1. FIRST: Parse the uploaded resume into a structured vault — extract every skill, experience entry, bullet point, metric, and education item into a tagged, organized format. This vault is your source of truth.\n"
        + voice_instruction + "\n"
        "3. Parse the job posting and classify requirements by signal strength (dealbreaker / strong / bonus)\n"
        "4. Analyze the CEO's pain — WHY they're hiring, not just what they want\n"
        "5. Generate the ideal candidate's resume for this role (the ghost)\n"
        "6. Map the vault's real experience onto the ghost resume's structure\n"
        "7. Generate gap report, cover letter, interview prep, ATS score\n\n"
        "CRITICAL RULES:\n"
        "- NEVER fabricate experience. Only reframe what the candidate actually has in their vault.\n"
        "- Mirror the posting's language without copying verbatim.\n"
        "- Quantify everything possible — pull metrics from the vault.\n"
        "- Address the CEO's pain in the professional summary and top bullets.\n"
        "- ALWAYS extract the candidate's full contact info (name, email, phone, location, linkedin) from the resume.\n"
        "- For the cover letter: search for the company's physical address and hiring manager name. Format as a proper business letter.\n"
        "- TECH STACK CORRECTION: If the resume lists technologies that contradict the actual work described (e.g. says 'Swift/Kotlin' but describes building a Flutter/Dart app), CORRECT IT in the vault and tailored output. The vault must reflect what was actually used, not what the resume mistakenly says. Cross-reference the work described against the tech listed.\n"
        "- EDUCATION IS MANDATORY: The tailored_resume MUST always include the education field with all degrees, institutions, and dates from the resume. NEVER omit education — ATS systems filter resumes without it.\n"
        "- SKILLS MUST INCLUDE HARD SKILLS: The skills array must contain SPECIFIC tools, technologies, platforms, and systems mentioned in the job posting — not just soft skill phrases. Extract every named tool/platform/system from the posting (CRM software, ticketing systems, programming languages, specific platforms) and include the ones the candidate has or can honestly claim. Soft skills can be included but must NOT be the majority. Aim for 60%+ hard/technical skills and max 40% soft skills.\n"
        "- CONTACT IN TAILORED RESUME: The contact field at the top level of the JSON response MUST be populated with name, email, phone, location from the resume. This is critical — documents cannot be generated without it.\n"
        + voice_consistency + "\n\n"
        """COVER LETTER HOOK PHILOSOPHY:
The hook must be COMPANY-SITUATION-CENTRIC, not candidate-centric.

BAD hooks (candidate-centric — recruiters are starting to pattern-match these):
- "I built X from a Y" (humble brag opener)
- "Most developers do X, but I do Y" (contrast template)
- "With N years of experience in X..." (resume summary)

GOOD hooks (company-situation-centric — makes the recruiter think about THEIR problem):
- Reference a specific challenge or opportunity the company is facing that most applicants wouldn't notice
- Frame an insight about the company's situation that the next sentence proves the candidate understands deeply
- Make the recruiter think "how does this person know about that?" — then the letter answers it

The hook should make the recruiter forget they're reading a cover letter and start thinking about their own problem. The candidate enters as the person who already understands it.

Write the full letter BODY first, then re-read it and REPLACE the first line with a company-situation-centric hook. The contact header, date, and recipient info go in separate fields — NOT in full_text.

SKILL TRANSLATION PHILOSOPHY (for the gap_report.skill_translations field):
Think of yourself as a lawyer presenting your client in the best possible light — you stretch the truth professionally, you never fabricate it. If the candidate mentions non-traditional experience (gaming, content creation, hobbies, volunteering, unconventional projects), identify real hard/soft skills those activities demonstrated and translate them into professional language.

Examples of valid skill translation:
- "Led a 40-person WoW raid guild" -> "Coordinated and led 40+ person team in high-pressure strategic operations"
- "Streamed on Twitch to 500 nightly viewers" -> "Built and engaged an audience of 500+ daily through consistent content production"
- "Modded a game server for 3 years" -> "Administered online community platforms, handling technical support and user management"

Rules for skill translation:
- The SOURCE activity must be real (from the resume or implied by context)
- The TRANSLATION must describe real skills actually demonstrated
- NEVER invent job titles, company names, or paid employment
- ALWAYS include a context_note explaining where/how to honestly frame it (Personal Projects section, interview discussion, etc.)
- If the candidate's resume has no gaps that need closing this way, return an empty skill_translations array

Respond ONLY with valid JSON (no markdown fences):
{
  "company": "string",
  "role": "string",
  "location": "string",
  "ceo_pain": "string",
  "pain_category": "string",
  "ats_score": number,
  "viability_score": number,
  "recommendation": "strong_apply | apply_with_strategy | risky_apply | skip",
  "red_flags": [{"flag": "string", "severity": "string"}],
  "green_flags": [{"flag": "string"}],
  "contact": {
    "name": "string (full name from resume)",
    "email": "string (from resume)",
    "phone": "string (from resume)",
    "location": "string (city, state from resume)",
    "linkedin": "string or null (from resume if present)"
  },
  "voice_print": {
    "sentence_style": "string (short_punchy | flowing | mixed)",
    "formality": "string (casual | professional | corporate)",
    "vocabulary_level": "string (plain | technical | mixed)",
    "leads_with": "string (results | context | action)",
    "personality_notes": "string (brief description of their natural writing voice)"
  },
  "vault": {
    "skills": [{"name": "string", "proficiency": "string", "tags": ["string"]}],
    "experience": [{"company": "string", "role": "string", "dates": "string", "bullets": [{"text": "string", "tags": ["string"], "metrics": "string or null"}]}],
    "education": [{"institution": "string", "degree": "string", "dates": "string"}]
  },
  "tailored_resume": {
    "summary": "string",
    "sections": [{"name": "string", "entries": [{"title": "string", "company": "string", "dates": "string", "bullets": ["string"]}]}],
    "skills": ["string"],
    "education": [{"degree": "string", "institution": "string", "dates": "string"}]
  },
  "cover_letter": {
    "recipient_name": "string (hiring manager name if found, otherwise 'Hiring Manager')",
    "recipient_title": "string or null",
    "company_name": "string",
    "company_address": "string (full street address of company, searched via web)",
    "hook_line": "string (COMPANY-SITUATION-CENTRIC — about their problem, not the candidate's achievement)",
    "original_first_line": "string",
    "full_text": "string (BODY ONLY — the 3 paragraphs, no header/date/address, no sign-off name)",
    "sign_off": "string (e.g. 'Sincerely,' or 'Best regards,')"
  },
  "gap_report": {
    "critical_gaps": [{"gap": "string", "strategy": "string"}],
    "advantages": [{"strength": "string", "pitch": "string"}],
    "gap_closers": [{"gap": "string", "tier": "string", "action": "string", "time": "string"}],
    "skill_translations": [{"source_activity": "string", "professional_translation": "string", "skills_demonstrated": ["string"], "context_note": "string"}]
  },
  "interview_prep": {
    "two_min_pitch": "string",
    "gap_questions": [{"question": "string", "answer": "string"}],
    "behavioral_stars": [{"question": "string", "situation": "string", "task": "string", "action": "string", "result": "string"}],
    "technical_questions": [{"question": "string", "answer": "string"}],
    "questions_to_ask": [{"question": "string", "why": "string"}],
    "salary": {"floor": "string", "target": "string", "stretch": "string"}
  }
}"""
    )

    user_message = f"RESUME:\n{request.resume_text}\n\nJOB POSTING:\n{job_text}"

    try:
        raw = await call_claude(system_prompt, user_message, use_search=True)
        result = parse_json_response(raw)
    except Exception as e:
        raise HTTPException(500, f"Pipeline error: {str(e)}")

    # ---- POST-PROCESSING VALIDATION ----
    # Ensure contact exists
    if not result.get("contact") or not result.get("contact", {}).get("name"):
        # Try to extract basic contact from resume text
        lines = request.resume_text.strip().split("\n")
        fallback_contact = {"name": "", "email": "", "phone": "", "location": "", "linkedin": ""}
        for line in lines[:10]:  # Contact info is usually in first 10 lines
            line = line.strip()
            if "@" in line and "." in line and not fallback_contact["email"]:
                # Likely email
                # re already imported at top
                email_match = re.search(r'[\w.+-]+@[\w-]+\.[\w.]+', line)
                if email_match:
                    fallback_contact["email"] = email_match.group()
            if re.search(r'\d{3}[\s.-]?\d{3}[\s.-]?\d{4}', line) and not fallback_contact["phone"]:
                phone_match = re.search(r'[\d\s.()+\-]{10,}', line)
                if phone_match:
                    fallback_contact["phone"] = phone_match.group().strip()
        if not fallback_contact["name"] and lines:
            # First non-empty line is usually the name
            for line in lines[:5]:
                clean = line.strip()
                if clean and len(clean) < 60 and "@" not in clean and not any(c.isdigit() for c in clean[:3]):
                    fallback_contact["name"] = clean
                    break
        if not result.get("contact"):
            result["contact"] = fallback_contact
        else:
            for k, v in fallback_contact.items():
                if v and not result["contact"].get(k):
                    result["contact"][k] = v

    # Ensure education exists in tailored_resume
    tr = result.get("tailored_resume", {})
    if not tr.get("education") or len(tr.get("education", [])) == 0:
        vault_edu = result.get("vault", {}).get("education", [])
        if vault_edu:
            tr["education"] = vault_edu
            result["tailored_resume"] = tr

    # Ensure skills has hard skills (warn if all soft)
    skills = tr.get("skills", [])
    soft_indicators = ["management", "collaboration", "communication", "leadership", "teamwork",
                       "problem-solving", "adaptability", "proactive", "customer service", "engagement"]
    if skills:
        soft_count = sum(1 for s in skills if any(ind in s.lower() for ind in soft_indicators))
        if soft_count > len(skills) * 0.6:
            result["_warning_soft_skills"] = True

    # Save session
    session_path = OUTPUT_DIR / f"{session_id}.json"
    result["session_id"] = session_id
    result["timestamp"] = datetime.now().isoformat()
    with open(session_path, "w") as f:
        json.dump(result, f, indent=2)

    return result


@app.get("/api/session/{session_id}")
async def get_session(session_id: str):
    """Retrieve a previous session's results."""
    session_path = OUTPUT_DIR / f"{session_id}.json"
    if not session_path.exists():
        raise HTTPException(404, "Session not found")
    with open(session_path) as f:
        return json.load(f)


class GapTranslateRequest(BaseModel):
    gap: str
    gap_strategy: str = ""
    vault: dict = {}
    role: str = ""
    company: str = ""


@app.post("/api/gap-translate")
async def gap_translate(request: GapTranslateRequest):
    """Look at the candidate's existing vault and propose honest translations
    for a specific gap, using only real experience they already have."""
    if not ANTHROPIC_KEY:
        raise HTTPException(500, "Anthropic API key not configured")

    system_prompt = """You are an ethical career advisor helping a candidate address a specific gap
in their resume using ONLY experience they actually have.

You will receive:
- A specific gap/missing skill the candidate has
- Their vault of real experience
- The role/company they're applying to

Your job: find any real experience in the vault that could honestly translate into the missing skill.
Think like a lawyer — present the candidate's real activities in their most favorable professional light,
but NEVER invent experiences they don't have.

RULES:
- Only use activities actually present in the vault
- Frame real experience using the vocabulary of the gap
- If NO honest translation is possible, return empty translations and say so — do not stretch
- Every translation needs an honest context_note about where it belongs

Respond ONLY with valid JSON:
{
  "can_be_bridged": true/false,
  "reasoning": "string (honest assessment of whether their real experience can address this gap)",
  "translations": [
    {
      "source_activity": "string (real thing from their vault)",
      "translated_bullet": "string (a bullet point ready to add to the resume, using the gap's vocabulary)",
      "section_to_add_to": "string (which resume section this should go in, e.g. 'Professional Experience', 'Projects', or 'Skills')",
      "context_note": "string (honest framing guidance)"
    }
  ]
}"""

    user_message = f"""GAP: {request.gap}
STRATEGY SUGGESTED: {request.gap_strategy}
ROLE/COMPANY: {request.role} at {request.company}

CANDIDATE'S VAULT (only use what's here):
{json.dumps(request.vault, indent=2)}"""

    try:
        raw = await call_claude(system_prompt, user_message)
        result = parse_json_response(raw)
        return result
    except Exception as e:
        raise HTTPException(500, f"Translation failed: {str(e)}")


class QuestionnaireStartRequest(BaseModel):
    gap: str
    gap_strategy: str = ""
    role: str = ""
    company: str = ""


@app.post("/api/gap-questionnaire/start")
async def gap_questionnaire_start(request: QuestionnaireStartRequest):
    """Start an interactive lawyer-style questionnaire for a specific gap.
    Claude generates the first question. Subsequent Q&A goes through /continue."""
    if not ANTHROPIC_KEY:
        raise HTTPException(500, "Anthropic API key not configured")

    system_prompt = """You are an ethical career advisor conducting a lawyer-style discovery
interview to uncover whether a candidate has REAL experience that addresses a specific gap.

Think like a lawyer interviewing a client: ask focused questions that might surface relevant
experience the candidate didn't think to include. Never suggest fabrication.

Start by asking ONE specific, open-ended question designed to surface real evidence of the
missing skill. The question should be practical and concrete, not abstract.

Example bad question: "Do you have any sales experience?"
Example good question: "Have you ever convinced someone to try something they were initially hesitant about — whether in a work, volunteer, or personal context? Walk me through what happened."

Respond ONLY with valid JSON:
{
  "question": "string (first question to ask)",
  "purpose": "string (what you're trying to uncover with this question)"
}"""

    user_message = f"""GAP TO EXPLORE: {request.gap}
CONTEXT: Applying for {request.role} at {request.company}
STRATEGY: {request.gap_strategy}

Generate the first discovery question."""

    try:
        raw = await call_claude(system_prompt, user_message)
        return parse_json_response(raw)
    except Exception as e:
        raise HTTPException(500, f"Questionnaire start failed: {str(e)}")


class QuestionnaireContinueRequest(BaseModel):
    gap: str
    role: str = ""
    company: str = ""
    conversation: list  # list of {question, answer} pairs so far
    question_count: int = 0  # how many questions asked so far


@app.post("/api/gap-questionnaire/continue")
async def gap_questionnaire_continue(request: QuestionnaireContinueRequest):
    """Continue the questionnaire. Claude either asks another question or concludes.
    After 3-5 questions, Claude must decide: can we honestly bridge this gap or not?"""
    if not ANTHROPIC_KEY:
        raise HTTPException(500, "Anthropic API key not configured")

    should_conclude = request.question_count >= 3

    system_prompt = f"""You are an ethical career advisor continuing a lawyer-style discovery
interview. Your job: uncover real experience that addresses a gap, OR honestly conclude it
can't be bridged.

RULES:
- You have asked {request.question_count} questions already.
- {'MUST CONCLUDE NOW: Produce final assessment.' if should_conclude else 'Either ask ONE more focused follow-up question (if answers suggest more to uncover) OR conclude if you have enough.'}
- If answers are thin, vague, or unrelated, be honest: this gap CANNOT be bridged from real experience.
- If answers reveal real relevant experience, produce resume bullets using ONLY what was said.

When concluding, respond with either:

A) If the candidate DOES have real relevant experience:
{{
  "action": "conclude_with_bullets",
  "can_be_bridged": true,
  "assessment": "string (honest summary of what real experience emerged)",
  "translations": [
    {{
      "source_activity": "string (what they actually described)",
      "translated_bullet": "string (resume bullet using ONLY what they said)",
      "section_to_add_to": "string (Professional Experience | Projects | Skills)",
      "context_note": "string (how to honestly frame it)"
    }}
  ]
}}

B) If the candidate does NOT have real relevant experience:
{{
  "action": "conclude_cannot_bridge",
  "can_be_bridged": false,
  "assessment": "string (honest explanation of why this gap can't be bridged from their real experience)",
  "recommendation": "string (what they should do instead — skip the role, pursue a gap closer action, or address in cover letter)"
}}

C) If you need ONE more question (only if question_count < 3):
{{
  "action": "ask_question",
  "question": "string",
  "purpose": "string"
}}"""

    conv_text = "\n\n".join([f"Q{i+1}: {qa['question']}\nA{i+1}: {qa['answer']}" for i, qa in enumerate(request.conversation)])

    user_message = f"""GAP: {request.gap}
ROLE/COMPANY: {request.role} at {request.company}

CONVERSATION SO FAR:
{conv_text}

Based on these answers, {'produce your final assessment' if should_conclude else 'decide: ask one more question, or conclude'}."""

    try:
        raw = await call_claude(system_prompt, user_message)
        return parse_json_response(raw)
    except Exception as e:
        raise HTTPException(500, f"Questionnaire continue failed: {str(e)}")


class DocumentRequest(BaseModel):
    type: str  # "resume" or "cover_letter"
    format: str  # "pdf" or "docx"
    data: dict  # the tailored_resume or cover_letter object from results
    company: str = "Company"
    role: str = "Role"


def _build_resume_docx(data: dict, contact: dict, filepath: str):
    """Generate a professional DOCX resume with contact header."""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def add_border(p):
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '1', qn('w:color'): '999999'})
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Name header
    if contact.get("name"):
        np = doc.add_paragraph()
        np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = np.add_run(contact["name"])
        nr.bold = True
        nr.font.size = Pt(18)
        nr.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        np.space_after = Pt(2)

    # Contact line
    contact_parts = [contact.get(k) for k in ["location", "phone", "email", "linkedin"] if contact.get(k)]
    if contact_parts:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(" | ".join(contact_parts))
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        cp.space_after = Pt(6)

    # Divider
    dp = doc.add_paragraph()
    dp.space_before = Pt(0)
    dp.space_after = Pt(4)
    pPr = dp._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {qn('w:val'): 'single', qn('w:sz'): '6', qn('w:space'): '1', qn('w:color'): '333333'})
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Summary
    summary = data.get("summary", "")
    if summary:
        p = doc.add_paragraph()
        run = p.add_run("PROFESSIONAL SUMMARY")
        run.bold = True
        run.font.size = Pt(11)
        add_border(p)
        sp = doc.add_paragraph(summary)
        sp.space_after = Pt(6)

    # Sections
    for section in data.get("sections", []):
        p = doc.add_paragraph()
        run = p.add_run((section.get("name", "")).upper())
        run.bold = True
        run.font.size = Pt(11)
        p.space_before = Pt(8)
        add_border(p)

        for entry in section.get("entries", []):
            ep = doc.add_paragraph()
            run_title = ep.add_run(entry.get("title", ""))
            run_title.bold = True
            run_title.font.size = Pt(10.5)
            if entry.get("company"):
                ep.add_run(f" — {entry['company']}")
            if entry.get("dates"):
                dr = ep.add_run(f"  |  {entry['dates']}")
                dr.font.size = Pt(9.5)
                dr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            ep.space_after = Pt(2)

            for bullet in entry.get("bullets", []):
                bp = doc.add_paragraph(style='List Bullet')
                bp.text = bullet
                bp.paragraph_format.left_indent = Inches(0.3)
                bp.paragraph_format.space_after = Pt(1)

    # Skills
    skills = data.get("skills", [])
    if skills:
        p = doc.add_paragraph()
        run = p.add_run("SKILLS")
        run.bold = True
        run.font.size = Pt(11)
        p.space_before = Pt(8)
        add_border(p)
        doc.add_paragraph(" | ".join(skills))

    # Education
    education = data.get("education", [])
    if education:
        p = doc.add_paragraph()
        run = p.add_run("EDUCATION")
        run.bold = True
        run.font.size = Pt(11)
        p.space_before = Pt(8)
        add_border(p)
        for edu in education:
            ep = doc.add_paragraph()
            er = ep.add_run(f"{edu.get('degree', '')}")
            er.bold = True
            ep.add_run(f" — {edu.get('institution', '')}  |  {edu.get('dates', '')}")

    doc.save(filepath)


def _build_resume_pdf(data: dict, contact: dict, filepath: str):
    """Generate a professional PDF resume with contact header."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    doc = SimpleDocTemplate(filepath, pagesize=letter,
                            topMargin=0.5*inch, bottomMargin=0.5*inch,
                            leftMargin=0.65*inch, rightMargin=0.65*inch)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='ResName', fontName='Helvetica-Bold', fontSize=16, alignment=1, spaceAfter=2, textColor=HexColor('#1a1a1a')))
    styles.add(ParagraphStyle(name='ResCont', fontName='Helvetica', fontSize=8.5, alignment=1, spaceAfter=4, textColor=HexColor('#666666')))
    styles.add(ParagraphStyle(name='SHead', fontName='Helvetica-Bold', fontSize=10, spaceBefore=10, spaceAfter=3, textColor=HexColor('#1a1a1a')))
    styles.add(ParagraphStyle(name='EntryTitle', fontName='Helvetica-Bold', fontSize=10, spaceAfter=1, textColor=HexColor('#333333')))
    styles.add(ParagraphStyle(name='EntryDate', fontName='Helvetica', fontSize=8.5, textColor=HexColor('#666666'), spaceAfter=2))
    styles.add(ParagraphStyle(name='ResBullet', fontName='Helvetica', fontSize=9.5, leftIndent=14, spaceAfter=2, textColor=HexColor('#333333')))
    styles.add(ParagraphStyle(name='Sum', fontName='Helvetica', fontSize=9.5, spaceAfter=6, textColor=HexColor('#333333')))
    styles.add(ParagraphStyle(name='Skl', fontName='Helvetica', fontSize=9.5, spaceAfter=4, textColor=HexColor('#333333')))

    story = []
    hr_thick = lambda: HRFlowable(width="100%", thickness=0.5, color=HexColor('#333333'), spaceAfter=6, spaceBefore=4)
    hr_thin = lambda: HRFlowable(width="100%", thickness=0.3, color=HexColor('#999999'), spaceAfter=4)

    # Contact header
    if contact.get("name"):
        story.append(Paragraph(contact["name"], styles['ResName']))
    contact_parts = [contact.get(k) for k in ["location", "phone", "email", "linkedin"] if contact.get(k)]
    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), styles['ResCont']))
    story.append(hr_thick())

    # Summary
    summary = data.get("summary", "")
    if summary:
        story.append(Paragraph("PROFESSIONAL SUMMARY", styles['SHead']))
        story.append(hr_thin())
        story.append(Paragraph(summary, styles['Sum']))

    for section in data.get("sections", []):
        story.append(Paragraph((section.get("name", "")).upper(), styles['SHead']))
        story.append(hr_thin())
        for entry in section.get("entries", []):
            title = f"<b>{entry.get('title', '')}</b>"
            if entry.get("company"):
                title += f" — {entry['company']}"
            story.append(Paragraph(title, styles['EntryTitle']))
            if entry.get("dates"):
                story.append(Paragraph(entry["dates"], styles['EntryDate']))
            for bullet in entry.get("bullets", []):
                safe = bullet.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(f"• {safe}", styles['ResBullet']))
            story.append(Spacer(1, 4))

    skills = data.get("skills", [])
    if skills:
        story.append(Paragraph("SKILLS", styles['SHead']))
        story.append(hr_thin())
        story.append(Paragraph(" | ".join(skills), styles['Skl']))

    education = data.get("education", [])
    if education:
        story.append(Paragraph("EDUCATION", styles['SHead']))
        story.append(hr_thin())
        for edu in education:
            safe = f"<b>{edu.get('degree', '')}</b> — {edu.get('institution', '')}  |  {edu.get('dates', '')}"
            story.append(Paragraph(safe, styles['Sum']))

    doc.build(story)


def _build_cover_letter_docx(data: dict, contact: dict, filepath: str):
    """Generate cover letter as DOCX with full business letter format."""
    from docx.shared import Pt, Inches, RGBColor
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Sender header
    if contact.get("name"):
        np = doc.add_paragraph()
        nr = np.add_run(contact["name"])
        nr.bold = True
        nr.font.size = Pt(13)
    contact_parts = [contact.get(k) for k in ["location", "phone", "email"] if contact.get(k)]
    if contact_parts:
        cp = doc.add_paragraph(" | ".join(contact_parts))
        for run in cp.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Date
    doc.add_paragraph("")
    doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))

    # Recipient
    doc.add_paragraph("")
    recipient = data.get("recipient_name", "Hiring Manager")
    doc.add_paragraph(recipient)
    if data.get("company_name"):
        doc.add_paragraph(data["company_name"])
    if data.get("company_address"):
        doc.add_paragraph(data["company_address"])

    # Greeting
    doc.add_paragraph("")
    greeting_name = recipient if recipient != "Hiring Manager" else "Hiring Team"
    doc.add_paragraph(f"Dear {greeting_name},")

    # Body
    doc.add_paragraph("")
    full_text = data.get("full_text", "")
    for para in full_text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # Sign-off
    doc.add_paragraph("")
    sign_off = data.get("sign_off", "Sincerely,")
    doc.add_paragraph(sign_off)
    doc.add_paragraph("")
    if contact.get("name"):
        doc.add_paragraph(contact["name"])

    doc.save(filepath)


def _build_cover_letter_pdf(data: dict, contact: dict, filepath: str):
    """Generate cover letter as PDF with full business letter format."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    doc = SimpleDocTemplate(filepath, pagesize=letter,
                            topMargin=1*inch, bottomMargin=1*inch,
                            leftMargin=1*inch, rightMargin=1*inch)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='SenderName', fontName='Helvetica-Bold', fontSize=13, spaceAfter=2, textColor=HexColor('#1a1a1a')))
    styles.add(ParagraphStyle(name='SenderInfo', fontName='Helvetica', fontSize=9.5, spaceAfter=2, textColor=HexColor('#666666')))
    styles.add(ParagraphStyle(name='LetterBody', fontName='Helvetica', fontSize=11, spaceAfter=10, textColor=HexColor('#333333'), leading=16))
    styles.add(ParagraphStyle(name='LetterMeta', fontName='Helvetica', fontSize=11, spaceAfter=2, textColor=HexColor('#333333')))

    story = []

    # Sender
    if contact.get("name"):
        story.append(Paragraph(contact["name"], styles['SenderName']))
    contact_parts = [contact.get(k) for k in ["location", "phone", "email"] if contact.get(k)]
    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), styles['SenderInfo']))
    story.append(Spacer(1, 16))

    # Date
    story.append(Paragraph(datetime.now().strftime("%B %d, %Y"), styles['LetterMeta']))
    story.append(Spacer(1, 12))

    # Recipient
    recipient = data.get("recipient_name", "Hiring Manager")
    story.append(Paragraph(recipient, styles['LetterMeta']))
    if data.get("company_name"):
        story.append(Paragraph(data["company_name"], styles['LetterMeta']))
    if data.get("company_address"):
        story.append(Paragraph(data["company_address"], styles['LetterMeta']))
    story.append(Spacer(1, 16))

    # Greeting
    greeting_name = recipient if recipient != "Hiring Manager" else "Hiring Team"
    story.append(Paragraph(f"Dear {greeting_name},", styles['LetterBody']))
    story.append(Spacer(1, 8))

    # Body
    full_text = data.get("full_text", "")
    for para in full_text.split("\n\n"):
        if para.strip():
            safe = para.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, styles['LetterBody']))

    # Sign-off
    story.append(Spacer(1, 12))
    sign_off = data.get("sign_off", "Sincerely,")
    story.append(Paragraph(sign_off, styles['LetterBody']))
    story.append(Spacer(1, 24))
    if contact.get("name"):
        story.append(Paragraph(contact["name"], styles['LetterMeta']))

    doc.build(story)


@app.post("/api/generate-document")
async def generate_document(request: DocumentRequest):
    """Generate a PDF or DOCX from tailored resume or cover letter data."""
    safe_company = "".join(c for c in request.company if c.isalnum() or c in (' ', '-', '_')).strip().replace(" ", "_")
    safe_role = "".join(c for c in request.role if c.isalnum() or c in (' ', '-', '_')).strip().replace(" ", "_")

    ext = request.format.lower()
    if ext not in ("pdf", "docx"):
        raise HTTPException(400, "Format must be 'pdf' or 'docx'")

    prefix = "Resume" if request.type == "resume" else "CoverLetter"
    filename = f"{prefix}_{safe_company}_{safe_role}.{ext}"
    filepath = OUTPUT_DIR / filename

    # Extract contact — try _contact first, then look for contact in data
    contact = request.data.get("_contact", {})
    if not contact or not contact.get("name"):
        # Fallback: try top-level contact field
        contact = request.data.get("contact", contact or {})

    try:
        if request.type == "resume":
            if ext == "pdf":
                _build_resume_pdf(request.data, contact, str(filepath))
            else:
                _build_resume_docx(request.data, contact, str(filepath))
        elif request.type == "cover_letter":
            if ext == "pdf":
                _build_cover_letter_pdf(request.data, contact, str(filepath))
            else:
                _build_cover_letter_docx(request.data, contact, str(filepath))
        else:
            raise HTTPException(400, "Type must be 'resume' or 'cover_letter'")
    except Exception as e:
        raise HTTPException(500, f"Document generation failed: {str(e)}")

    return FileResponse(
        path=str(filepath),
        filename=filename,
        media_type="application/pdf" if ext == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ============================================================
# OUTCOME TRACKING
# ============================================================
OUTCOMES_PATH = OUTPUT_DIR / "outcomes.json"


class OutcomeReport(BaseModel):
    session_id: str
    company: str = ""
    role: str = ""
    outcome: str  # "got_interview" | "got_offer" | "rejected" | "no_response"
    notes: str = ""


def _load_outcomes():
    if OUTCOMES_PATH.exists():
        with open(OUTCOMES_PATH) as f:
            return json.load(f)
    return {"outcomes": [], "stats": {"total": 0, "interviews": 0, "offers": 0, "rejections": 0, "no_response": 0}}


def _save_outcomes(data):
    with open(OUTCOMES_PATH, "w") as f:
        json.dump(data, f, indent=2)


@app.post("/api/report-outcome")
async def report_outcome(report: OutcomeReport):
    """User reports what happened after applying."""
    outcomes = _load_outcomes()
    outcomes["outcomes"].append({
        "session_id": report.session_id,
        "company": report.company,
        "role": report.role,
        "outcome": report.outcome,
        "notes": report.notes,
        "timestamp": datetime.now().isoformat(),
    })
    outcomes["stats"]["total"] += 1
    if report.outcome == "got_interview":
        outcomes["stats"]["interviews"] += 1
    elif report.outcome == "got_offer":
        outcomes["stats"]["offers"] += 1
    elif report.outcome == "rejected":
        outcomes["stats"]["rejections"] += 1
    elif report.outcome == "no_response":
        outcomes["stats"]["no_response"] += 1
    _save_outcomes(outcomes)
    return {"status": "recorded", "stats": outcomes["stats"]}


@app.get("/api/stats")
async def get_stats():
    """Public aggregate stats for social proof on landing page."""
    outcomes = _load_outcomes()
    sessions_count = len(list(OUTPUT_DIR.glob("*.json"))) - 1  # minus outcomes.json
    stats = outcomes["stats"]
    interview_rate = round((stats["interviews"] + stats["offers"]) / max(stats["total"], 1) * 100)
    return {
        "total_resumes_tailored": max(sessions_count, 0),
        "outcomes_reported": stats["total"],
        "interview_rate": interview_rate,
        "interviews": stats["interviews"],
        "offers": stats["offers"],
    }


@app.get("/health")
async def health():
    return {
        "status": "ok",
        "anthropic_key_set": bool(ANTHROPIC_KEY),
        "perplexity_key_set": bool(PERPLEXITY_KEY),
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
